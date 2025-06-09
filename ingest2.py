from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_cohere import CohereEmbeddings
from langchain_chroma import Chroma
from uuid import uuid4
import os
import json 
import pdfplumber 
from docx import Document 
import win32com.client as win32
import pandas as pd

def process_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return 
    
    text = ""
    doc = Document(file_path)
    for para in doc.paragraphs:
        # Extract text from paragraphs
        text += para.text.strip()
    
    table_dfs = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
    
        # Gán header nếu đủ điều kiện, ngược lại dùng số cột để đặt tên cột
        if data:
            header = data[0]
            if all(header):  # Nếu tất cả cột đầu tiên có tiêu đề
                df = pd.DataFrame(data[1:], columns=header)
            else:
                col_count = len(header)
                columns = [f"Column {i+1}" for i in range(col_count)]
                df = pd.DataFrame(data, columns=columns)
            table_dfs.append(df)
            
    return text, table_dfs
    
def process_pdf(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return None, None

    all_text = ""
    all_tables = []

    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages):
            # Trích xuất text từ trang
            text = page.extract_text()
            if text:
                all_text += text + "\n"

            # Trích xuất bảng từ trang
            raw_tables = page.extract_tables()
            for table in raw_tables:
                if not table:
                    continue
                # Kiểm tra và làm sạch bảng
                cleaned_table = [[cell.strip() if cell else "" for cell in row] for row in table]
                if cleaned_table:
                    header = cleaned_table[0]
                    if all(header):
                        df = pd.DataFrame(cleaned_table[1:], columns=header)
                    else:
                        col_count = len(header)
                        columns = [f"Column {i+1}" for i in range(col_count)]
                        df = pd.DataFrame(cleaned_table, columns=columns)
                    all_tables.append(df)

    return all_text.strip(), all_tables

def convert_doc_to_docx(doc_path):
    """
    Chuyển đổi một file .doc thành .docx bằng cách sử dụng Microsoft Word.
    Yêu cầu: Chạy trên Windows và có cài đặt MS Word.
    """
    # Tạo một đường dẫn tuyệt đối để Word có thể tìm thấy file
    doc_path_abs = os.path.abspath(doc_path)
    docx_path_abs = doc_path_abs + 'x'

    # Kiểm tra nếu file .docx đã tồn tại thì không cần convert nữa
    if os.path.exists(docx_path_abs):
        print(f"File .docx đã tồn tại: {docx_path_abs}")
        return docx_path_abs

    try:
        # Khởi tạo ứng dụng Word
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False # Chạy ẩn

        # Mở file .doc
        doc = word.Documents.Open(doc_path_abs)
        
        # Lưu file sang định dạng .docx (wdFormatXMLDocument = 12)
        doc.SaveAs(docx_path_abs, FileFormat=12)
        doc.Close()
        word.Quit()
        
        print(f"Đã chuyển đổi thành công: {docx_path_abs}")
        return docx_path_abs
    except Exception as e:
        print(f"Lỗi khi chuyển đổi file {doc_path}: {e}")
        # Đảm bảo Word được đóng lại nếu có lỗi
        if 'word' in locals() and word is not None:
            word.Quit()
        return None

def process_doc(file_path):
    """
    Hàm xử lý file .doc bằng cách chuyển nó sang .docx rồi xử lý.
    """
    print(f"Phát hiện file .doc, đang tiến hành chuyển đổi: {file_path}")
    docx_path = convert_doc_to_docx(file_path)
    
    if docx_path:
        # After converting, process the .docx file
        print("--- Bắt đầu xử lý file .docx vừa được chuyển đổi ---")
        return process_docx(docx_path)
        

with open('data/document-list.json', 'r', encoding='utf-8') as f: 
    data = json.load(f)
    
documents = data["contents"]
print(len(documents))

for doc in documents:
    print(f"--- {doc['mediaTitle']} ---")
    for attach in doc.get("attachmentList", []):
        file_path = attach["url"]
        file_title = attach["title"]
        file_path = file_path.replace("/home/ctct_hdqt_owner/qltt_web_8882/Upload/QLTT", "data/QLTT_1")
        print(f"Tệp: {file_title} | Đường dẫn: {file_path}")
        if file_path.lower().endswith('.pdf') == True: 
            text, tables = process_pdf(file_path)
        elif file_path.lower().endswith('.docx') == True: 
            text, tables = process_docx(file_path)
        elif file_path.lower().endswith('.doc') == True:
            text, tables = process_doc(file_path)
    break
